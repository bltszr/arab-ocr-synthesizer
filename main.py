#!/usr/bin/env python

from arabic_reshaper import ArabicReshaper
from bidi.algorithm import get_display

from PIL import Image, ImageDraw, ImageFont

import docx
from docx.shared import Length, Inches, Pt, Mm, Emu

import math

# from pdfminer.high_level import extract_pages
# from pdfminer.layout import LTTextContainer, LTChar, LTTextLine 

import argparse
import os
import json
import random
import glob
import sys
import shutil

import time
from datetime import datetime

from data_report import check_font_on_text, path2font
from util import real_preprocess, render_preprocess

from functools import partial

tatwil = 'ـ'

to_px = lambda length, dpi : int((length.inches * dpi))
from_px = lambda pxs, dpi : Inches(pxs / dpi)

line_spacing_rules = {
  docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE: 1.5,
  docx.enum.text.WD_LINE_SPACING.SINGLE: 1,
  docx.enum.text.WD_LINE_SPACING.AT_LEAST: 1,
  docx.enum.text.WD_LINE_SPACING.DOUBLE: 2,
}

line_spacing_map = {
  "1.5": docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE,
  "single": docx.enum.text.WD_LINE_SPACING.SINGLE,
  "double": docx.enum.text.WD_LINE_SPACING.DOUBLE,
}

DEFAULT = {
  # "spacing": Inches(0.5),
  # "spacing rule": 1,
  # "font": "Calibri",
  # "font size": Pt(12),
  # "page height": Mm(297),
  # "page width": Mm(210),
  # "left margin": Inches(1),
  # "right margin": Inches(1),
  # "top margin": Inches(1),
  # "bottom margin": Inches(1),
}

backrounds = []

to_project_dir = lambda path, output_dir : os.path.join(output_dir, f"{os.path.basename(path)}.{datetime.now().timestamp()}.d")

def save(img, bboxes, dest_folder, page_number, verbose=False):
  img_filename = os.path.join(dest_folder,  f'{page_number}.png')
  json_filename = os.path.join(dest_folder,  f'{page_number}.json')
  img.save(img_filename)
  with open(json_filename, 'w') as f:
    json.dump(bboxes, f, ensure_ascii=False)
    if verbose:
      print(f'[INFO] Saved \'{img_filename}\'')


def create_page(page_width, page_height, bg=False):
  global backgrounds
  if bg:
    background = random.choice(backgrounds)
    img = Image.open(background)\
               .resize((page_width, page_height))\
               .convert("RGBA")
  else:
    img = Image.new("RGBA",
                    tuple(map(int,
                              (page_width,
                               page_height))),
                    color="white")
  draw = ImageDraw.Draw(img)
  return img, draw, []

# from https://stackoverflow.com/a/67203353
def get_wrapped_text(text: str,
                     font: ImageFont.ImageFont,
                     line_length: int,
                     delim=' '):
  lines = ['']
  for word in text.split(delim):
    line = f'{lines[-1]}{delim}{word}'.strip()
    if font.getlength(line, direction="rtl") <= line_length:
      lines[-1] = line
    else:
      lines.append(word)
  return lines


def get_font_size(style):
    if style.font.size == None:
        if style.base_style != None:
            return get_font_size(style.base_style)
        else:
            return None
    else:
        return style.font.size

def get_font_name(style):
  if style.font.name == None:
    if style.base_style != None:
      return get_font_size(style.base_style)
    else:
      return None
  else:
    return style.font.name
  
def get_font(run, font_dict, dpi):
  # Get the font size of the run
  try:
    run_font_size = get_font_size(run.style)
  except (TypeError, AttributeError):
    run_font_size = DEFAULT["font-size"]

  try:
    run_font_name = get_font_name(to_px(run.style, dpi))
  except (TypeError, AttributeError):
    run_font_name = DEFAULT["font"]
    
  return ImageFont.truetype(font_dict[run_font_name],
                            to_px(run_font_size, dpi))

def get_spacing(paragraph):
  try:
    if type(paragraph.paragraph_format.line_spacing) == float:
      try:
        spacing = from_px(get_font_size(paragraph.style) * paragraph.paragraph_format.line_spacing, dpi)
      except (TypeError, AttributeError, IndexError) as e:
        spacing = Emu(DEFAULT["spacing"] * DEFAULT["spacing-rule"])
    else:
      spacing = paragraph.paragraph_format.line_spacing * line_spacing_rules[
        paragraph.paragraph_format.line_spacing_rules
      ]
  except (TypeError, AttributeError) as e:
    spacing = Emu(DEFAULT["spacing"] * DEFAULT["spacing-rule"])
  return spacing

def get_indents(paragraph):
  right_indent = paragraph.paragraph_format.right_indent
  left_indent = paragraph.paragraph_format.left_indent
  if left_indent == None:
    left_indent = Inches(0)
  if right_indent == None:
    right_indent = Inches(0)
  return right_indent, left_indent

def process_txt(args, font):
  font_for_checking = path2font(args.font)
  dest_folder = to_project_dir(args.path, args.output_dir)
  os.makedirs(dest_folder, exist_ok=True)
  page_width, page_height,\
    left_margin, right_margin,\
    top_margin, bottom_margin = map(partial(to_px, dpi=args.dpi),
                                    (DEFAULT['page-width'],
                                     DEFAULT['page-height'],
                                     DEFAULT['left-margin'],
                                     DEFAULT['right-margin'],
                                     DEFAULT['top-margin'],
                                     DEFAULT['bottom-margin']))
  page_number = 0
  # Iterate over each paragraph in the document
  img, draw, bboxes = create_page(page_width, page_height, args.background)
  
  cum_spacing = top_margin
  # Extract the text and style of the paragraph
  par_spacing = to_px(Emu(DEFAULT['spacing'] \
                          * line_spacing_rules[DEFAULT['spacing-rule']]),
                      args.dpi)
  right_indent = 0
  left_indent = 0
  for j, run in enumerate(open(args.path, 'r')):
    if page_number > args.end_page:
      break
    if args.scriptio_continuo:
      run = run.replace(' ', "\u200c")
    text = get_wrapped_text(run,
                            font,
                            page_width \
                            - right_margin \
                            - left_margin \
                            - right_indent \
                            - left_indent,
                            delim="\u200c" if args.scriptio_continuo else' ')
    
    for k, line in enumerate(text):
      # if line 
      # middle part = page_width - right_indent - left_indent - right_margin - left_margin
      # x = left_margin + left_indent + (middle part - length of line)
      # x = page_width - right_indent - right_margin - length of line
      x = page_width \
        - right_margin - right_indent \
        - font.getlength(line)
      if cum_spacing + font.size > page_height - bottom_margin:
        if page_number == 0:
          page_number += 1
        save(img, bboxes, dest_folder, page_number, args.verbose)
        img, draw, bboxes = create_page(page_width, page_height, args.background)
        cum_spacing = top_margin
        page_number += 1
      if page_number > args.end_page:
        break
      y = cum_spacing
      if y > page_height:
        if args.warn:
          print(f"[WARN] Ran off page at line {line}\n")
        continue
      
      line = real_preprocess(line)
      line_to_render = render_preprocess(line)
      # check first
      missing_chars = check_font_on_text(font_for_checking, line)
      if len(missing_chars) > 0:
        if args.warn:
          print(f'[WARN] Font "{args.font}" missing these characters {missing_chars} present in file \"{args.path}\". Skipping line...')
        continue
      if len(line) == 0:
        if args.warn:
          print(f"[WARN] Empty line. Skipping...")
        continue
      alpha = int(random.uniform(args.min_alpha, args.max_alpha) * 255)
      txt_im = Image.new('RGBA', img.size,
                         (255,255,255,0))
      txt_d = ImageDraw.Draw(txt_im)
      txt_d.text((x, y), line_to_render,
                 font=font,
                 fill=(0, 0, 0, alpha),
                 direction="rtl")
      bbox = draw.textbbox((x, y), line_to_render,
                           font=font,
                           direction="rtl")
      img = Image.alpha_composite(img, txt_im)
      bboxes.append({"text": line,
                     "bbox": bbox})
      cum_spacing += par_spacing
      # save(img, bboxes, dest_folder, page_number)
  if page_number > 0:
    with open(os.path.join(dest_folder, 'config.json'), 'w') as f:
      json.dump(vars(args), f)
  else:
    shutil.rmtree(dest_folder)

def process_chars(args, font):
  # check first
  missing_chars = check_font_on_text(args.font, args.path)
  if len(missing_chars) > 0:
    raise ValueError(f'Font "{args.font}" missing these characters {missing_chars} present in file \"{args.path}\"')
  
  
  dest_folder = to_project_dir(args.path, args.output_dir)
  os.makedirs(dest_folder, exist_ok=True)
  page_width, page_height,\
    left_margin, right_margin,\
    top_margin, bottom_margin = map(partial(to_px, dpi=args.dpi),
                                    (DEFAULT['page-width'],
                                     DEFAULT['page-height'],
                                     DEFAULT['left-margin'],
                                     DEFAULT['right-margin'],
                                     DEFAULT['top-margin'],
                                     DEFAULT['bottom-margin']))
  page_number = 1
  # Iterate over each paragraph in the document
  img, draw, bboxes = create_page(page_width, page_height, args.background)
  
  cum_spacing = top_margin
  # Extract the text and style of the paragraph
  par_spacing = to_px(Emu(DEFAULT['spacing'] \
                          * line_spacing_rules[DEFAULT['spacing-rule']]).
                      args.dpi)
  right_indent = 0
  left_indent = 0
  configuration = {
    'delete_harakat': False,
    'support_ligatures': True,
    'RIAL SIGN': True,  # Replace ر ي ا ل with ﷼
  }
  reshaper = ArabicReshaper(configuration=configuration)
  for j, run in enumerate(open(args.path, 'r')):
    if page_number > args.end_page:
      break
    text = get_wrapped_text(run,
                            font,
                            page_width \
                            - right_margin \
                            - left_margin \
                            - right_indent \
                            - left_indent)
    for k, line in enumerate(text):
      x = page_width \
        - right_margin - right_indent \
        - font.getlength(line)
      if cum_spacing + font.size > page_height - bottom_margin:
        # save(img, bboxes, dest_folder, page_number)
        img, draw, bboxes = create_page(page_width, page_height, args.background)
        cum_spacing = top_margin
        page_number += 1
      if page_number > args.end_page:
        break
      y = cum_spacing
      if args.warn and y > page_height:
        print(f"[WARN] Ran off page at line {line}\n")
        # line = reshaper.reshape(get_display(line))
      
      line = u"\u200f" + line.replace('﴾', '(').replace('﴿', ')')
      
      alpha = int(random.uniform(args.min_alpha, args.max_alpha) * 255)
      txt_im = Image.new('RGBA', img.size,
                         (255,255,255,0))
      txt_d = ImageDraw.Draw(txt_im)
      txt_d.text((x, y), line,
                 font=font,
                 fill=(0, 0, 0, alpha),
                 direction="rtl")
      left, top, right, bottom = draw.textbbox((x, y), line,
                                               font=font,
                                               direction="rtl")
      if tatwil in line:
        left_t, _, right_t, _ = draw.textbbox((x, y), tatwil,
                                              font=font,
                                              direction="rtl")
        tatwil_width = right_t - left_t
        if line.startswith(tatwil):
          right -= tatwil_width
        if line.endswith(tatwil):
          left += tatwil_width
      char_path = os.path.join(dest_folder,
                               line.replace(tatwil, '').replace('\u200f', ''))
      os.makedirs(char_path, exist_ok=True)
      img = Image.alpha_composite(img, txt_im)
      img.crop((left, top, right, bottom))\
         .save(os.path.join(char_path,
                            f'{datetime.now()}.png'))
      # bboxes.append({"text": line,
      #                "bbox": (left, top, right, bottom)})
      cum_spacing += par_spacing
      # save(img, bboxes, dest_folder, page_number)
  with open(os.path.join(dest_folder, 'config.json'), 'w') as f:
    json.dump(vars(args), f)

def process_doc(args, font_dict):
  dest_folder = to_project_dir(args.path, args.output_dir)
  doc = docx.Document(doc_path)
  os.makedirs(dest_folder, exist_ok=True)

  section = doc.sections[0]
  page_width, page_height,\
    left_margin, right_margin,\
    top_margin, bottom_margin = map(partial(to_px, dpi=args.dpi),
                                    (section.page_width,
                                     section.page_height,
                                     section.left_margin,
                                     section.right_margin,
                                     section.top_margin,
                                     section.bottom_margin))
  
  page_number = 1
  text_page_height = page_height - top_margin - bottom_margin
  # Iterate over each paragraph in the document
  img, draw, bboxes = create_page(page_width, page_height, bg)
  cum_spacing = top_margin
  for i, paragraph in enumerate(doc.paragraphs):
    if page_number > args.end_page:
      break 
    # Extract the text and style of the paragraph
    right_indent, left_indent = get_indents(paragraph)
    style = paragraph.style.name
    par_spacing = to_px(get_spacing(paragraph), args.dpi)
    for j, run in enumerate(paragraph.runs):
      
      if run._element.xpath('w:lastRenderedPageBreak'):
        save(img, bboxes, dest_folder, page_number)
        img, draw, bboxes = create_page(page_width, page_height)
        cum_spacing = top_margin
        page_number += 1
      if page_number > args.end_page:
        break
      font = get_font(run, font_dict)
      text = get_wrapped_text(run.text,
                              font,
                              page_width \
                              - right_margin \
                              - left_margin \
                              - right_indent \
                              - left_indent)
      for k, line in enumerate(text):
        if line == '':
          continue
        # middle part = page_width - right_indent - left_indent - right_margin - left_margin
        # x = left_margin + left_indent + (middle part - length of line)
        # x = page_width - right_indent - right_margin - length of line
        x = page_width \
          - right_margin - right_indent \
          - font.getlength(line)
        if cum_spacing + font.size > text_page_height:
          save(img, bboxes, dest_folder, page_number)
          img, draw, bboxes = create_page(page_width, page_height)
          cum_spacing = top_margin
          page_number += 1
        if page_number > args.end_page:
          break
        y = cum_spacing
        if y > page_height:
          if args.warn:
            print(f"[WARN] Ran off page at line {line}\n")
          draw.text((x, y), line,
                    font=font,
                    fill=(0, 0, 0),
                    direction="rtl")
          bbox = draw.textbbox((x, y), line,
                               font=font,
                               direction="rtl")
          bboxes.append({"text": line,
                         "bbox": bbox})
          cum_spacing += par_spacing


def process_doc(args, font_dict):
  dest_folder = to_project_dir(args.path, args.output_dir)
  doc = docx.Document(doc_path)
  os.makedirs(dest_folder, exist_ok=True)

  section = doc.sections[0]
  page_width, page_height,\
    left_margin, right_margin,\
    top_margin, bottom_margin = map(partial(to_px, dpi=args.dpi),
                                    (section.page_width,
                                     section.page_height,
                                     section.left_margin,
                                     section.right_margin,
                                     section.top_margin,
                                     section.bottom_margin))
  
  page_number = 1
  text_page_height = page_height - top_margin - bottom_margin
  # Iterate over each paragraph in the document
  img, draw, bboxes = create_page(page_width, page_height, bg)
  cum_spacing = top_margin
  for i, paragraph in enumerate(doc.paragraphs):
    if page_number > args.end_page:
      break 
    # Extract the text and style of the paragraph
    right_indent, left_indent = get_indents(paragraph)
    style = paragraph.style.name
    par_spacing = to_px(get_spacing(paragraph), args.dpi)
    for j, run in enumerate(paragraph.runs):
      
      if run._element.xpath('w:lastRenderedPageBreak'):
        save(img, bboxes, dest_folder, page_number)
        img, draw, bboxes = create_page(page_width, page_height)
        cum_spacing = top_margin
        page_number += 1
      if page_number > args.end_page:
        break
      font = get_font(run, font_dict)
      text = get_wrapped_text(run.text,
                              font,
                              page_width \
                              - right_margin \
                              - left_margin \
                              - right_indent \
                              - left_indent)
      for k, line in enumerate(text):
        if line == '':
          continue
        # middle part = page_width - right_indent - left_indent - right_margin - left_margin
        # x = left_margin + left_indent + (middle part - length of line)
        # x = page_width - right_indent - right_margin - length of line
        x = page_width \
          - right_margin - right_indent \
          - font.getlength(line)
        if cum_spacing + font.size > text_page_height:
          save(img, bboxes, dest_folder, page_number)
          img, draw, bboxes = create_page(page_width, page_height)
          cum_spacing = top_margin
          page_number += 1
        if page_number > args.end_page:
          break
        y = cum_spacing
        if y > page_height:
          if args.warn:
            print(f"[WARN] Ran off page at line {line}\n")
          draw.text((x, y), line,
                    font=font,
                    fill=(0, 0, 0),
                    direction="rtl")
          bbox = draw.textbbox((x, y), line,
                               font=font,
                               direction="rtl")
          bboxes.append({"text": line,
                         "bbox": bbox})
          cum_spacing += par_spacing
          
def main(args):
  global backgrounds
  backgrounds = glob.glob(os.path.join(args.background, 'out*.png'))
  if args.alpha:
    args.min_alpha = args.alpha
    args.max_alpha = args.alpha
  if args.margin:
    args.bottom_margin = args.margin
    args.top_margin = args.margin
    args.right_margin = args.margin
    args.left_margin = args.margin

  DEFAULT.update({"font": args.font,
                  "font-size": Pt(args.font_size),
                  "spacing": Inches(args.spacing),
                  "left-margin": Inches(args.left_margin),
                  "right-margin": Inches(args.right_margin),
                  "top-margin": Inches(args.top_margin),
                  "bottom-margin": Inches(args.bottom_margin),
                  "page-width": Mm(args.page_width),
                  "page-height": Mm(args.page_height),
                  "spacing-rule": line_spacing_map[args.spacing_rule]})
  
  with open('fonts') as fp:
    font_dict = json.load(fp)
  if args.verbose:
    args.warn = True
  if args.path.endswith(".docx"):
    process_doc(args, font_dict)
  elif args.path.endswith(".txt") or args.path.endswith('.chars'):
    if os.path.exists(args.font):
      fontpath = args.font
    else:
      fontpath = font_dict[args.font]
    assert os.path.exists(fontpath)
    font = ImageFont.truetype(fontpath, size=to_px(DEFAULT['font-size'], args.dpi))
    if args.path.endswith('.txt'):
      process_txt(args, font)
    elif args.path.endswith('.chars'):
      process_chars(args, font)
      
if __name__ == "__main__":
  parser = argparse.ArgumentParser(description="Synthesizer for OCR data")
  parser.add_argument("path", type=str, help="Path to docx or txt file")
  parser.add_argument("--start-page", type=int, help="Starter page to parse",default=0)
  parser.add_argument("--end-page", type=int,
                      help="Starter page to parse, will be ignored if file is txt", default=math.inf)
  parser.add_argument("--font", type=str,
                      help="Override default font",
                      default='Calibri')
  parser.add_argument("--font-size", type=int,
                      help="Override default font size (pt)", default=12)
  parser.add_argument("--spacing", type=float,
                      help="Override default spacing size (inches)", default=0.5)
  parser.add_argument("--spacing-rule", choices=['1.5', 'double', 'single'], default='single')
  
  parser.add_argument("--min-alpha", type=float, default=0.85,
                      help="Minimum alpha to choose for text")
  parser.add_argument("--max-alpha", type=float, default=1.0,
                      help="Maximum alpha to choose for text")
  parser.add_argument("--alpha", type=float, default=None,
                      help="Single alpha value for all text. Overrides min-alpha and max-alpha")

  parser.add_argument("--left-margin", type=float, default=1, help="Override left margin (inches)")
  parser.add_argument("--right-margin", type=float, default=1, help="Override right margin (inches)")
  parser.add_argument("--top-margin", type=float, default=1, help="Override top margin (inches)")
  parser.add_argument("--bottom-margin", type=float, default=1, help="Override bottom margin (inches)")
  parser.add_argument("--margin", type=float, default=None, help="Override all margins (inches)")
  
  parser.add_argument("--page-width", type=float, default=210, help="Override page width (mm)",)
  parser.add_argument("--page-height", type=float, default=297, help="Override page height (mm)")

  parser.add_argument("--dpi", type=int, default=200, help="Dots per inch")
  
  parser.add_argument("--background", type=str, help="Path to folder of backgrounds")

  parser.add_argument("--output-dir", "-o", type=str, help="Path to output directory", default='outputs')
  
  parser.add_argument("--warn", action="store_true", help="Emit warnings")
  parser.add_argument("--verbose", action="store_true", help="Set warnings and info to true")
  parser.add_argument('--scriptio-continuo', '-sc', action='store_true', help="Replace spaces with zero-width non-joiners")
  
  main(parser.parse_args())
